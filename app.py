import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import requests
import traceback

st.set_page_config(page_title="Generator SK BPS", layout="centered")

st.title("📝 Generator Dokumen SK BPS Kota Solok")
st.write("Aplikasi ini akan mengotomasi pengisian dokumen. Template diambil langsung dari GitHub.")

# --- KONFIGURASI GITHUB ---
# Pastikan link ini mengarah ke file DOK.docx versi terbaru Anda
GITHUB_RAW_URL = "https://raw.githubusercontent.com/alberanalafean22/testotomasidokumen/main/DOK.docx"

# 1. Input Data Umum
st.subheader("1. Informasi Umum Dokumen")
col1, col2 = st.columns(2)

with col1:
    nomor = st.text_input("Nomor Dokumen", placeholder="Contoh: 042.1 TAHUN 2026")
    tanggal = st.text_input("Tanggal Ditetapkan", placeholder="Contoh: 15 Januari 2026")
    petugas = st.text_input("Menetapkan...", placeholder="Contoh: Petugas Sensus Ekonomi 2026")

with col2:
    tentang = st.text_input("Judul", placeholder="Contoh: PETUGAS SENSUS EKONOMI 2026 [tulis dengan huruf besar")
    pelaksanaan = st.text_input("Pelaksanaan Kegiatan", placeholder="Contoh: Sensus Ekonomi 2026")

# 2. Input Data Tabel (Dinamis)
st.subheader("2. Data Petugas (Lampiran)")
default_data = pd.DataFrame({
    "Nama/Jabatan": ["Alber Analafean"],
    "NIP/Golongan": ["199801012025011001 / III/a"],
    "Posisi": ["Ketua Tim"],
    "Honor": ["500.000"]
})
edited_df = st.data_editor(default_data, num_rows="dynamic", use_container_width=True)

if "doc_ready" not in st.session_state:
    st.session_state.doc_ready = False
if "doc_data" not in st.session_state:
    st.session_state.doc_data = None
if "doc_name" not in st.session_state:
    st.session_state.doc_name = ""

# --- FUNGSI CUSTOM FONT & BORDER ---
def apply_bookman_font(run, size=11):
    """Menerapkan font Bookman Old Style ke dalam elemen Word"""
    run.font.name = 'Bookman Old Style'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Bookman Old Style')
    run.font.size = Pt(size)

def set_cell_text_with_font(cell, text):
    """Memasukkan teks ke sel tabel sambil mempertahankan font Bookman"""
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            apply_bookman_font(run, size=11)

def set_bottom_border(cell):
    """Menerapkan garis ganda hitam (double border) di bagian bawah sebuah sel"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Mencari tag batas (borders) menggunakan find() agar kompatibel dan bebas error
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
        
    # Memeriksa batas bawah
    bottom = tcBorders.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        tcBorders.append(bottom)
        
    # Mengatur ketebalan (sz=12 setara dengan 1.5 pt), tipe double, warna hitam (000000)
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')

# 3. Tombol Eksekusi
if st.button("Proses & Buat Dokumen", type="primary"):
    with st.spinner('Mengunduh template dari GitHub dan memproses dokumen...'):
        try:
            response = requests.get(GITHUB_RAW_URL)
            
            if response.status_code != 200:
                st.error(f"Gagal mengunduh template (Status: {response.status_code}). Pastikan link Raw GitHub benar.")
                st.session_state.doc_ready = False
            else:
                template_file = io.BytesIO(response.content)
                doc = Document(template_file)
                
                replacements = {
                    "{tentang}": tentang,
                    "{pelaksanaan}": pelaksanaan,
                    "{pelaksanan}": pelaksanaan,
                    "{petugas}": petugas,
                    "{tanggal}": tanggal,
                    "{nomor}": nomor
                }
                
                # Fungsi replace teks
                def replace_text_in_paragraphs(paragraphs):
                    for p in paragraphs:
                        original_text = p.text
                        changed = False
                        for key, value in replacements.items():
                            if key in original_text:
                                original_text = original_text.replace(key, value)
                                changed = True
                        
                        if changed:
                            p.text = original_text
                            for run in p.runs:
                                apply_bookman_font(run, size=11)

                replace_text_in_paragraphs(doc.paragraphs)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if "{nama}" not in cell.text.lower():
                                replace_text_in_paragraphs(cell.paragraphs)

                # Logik Tabel Lampiran
                target_table = None
                template_row_idx = -1
                
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        for cell in row.cells:
                            if "{nama}" in cell.text.lower() or "{nama}" in cell.text:
                                target_table = table
                                template_row_idx = i
                                break
                        if target_table: break
                    if target_table: break
                
                if target_table and template_row_idx != -1:
                    for index, row_data in edited_df.iterrows():
                        if index == 0:
                            # Jika orang pertama, gunakan baris asli template
                            target_row = target_table.rows[template_row_idx]
                        else:
                            # Jika orang ke 2, 3, dst, TAMBAHKAN BARIS KOSONG (sebagai 1 spasi antar baris)
                            spacer = target_table.add_row()
                            
                            # Kemudian, tambahkan baris baru untuk data orang tersebut
                            target_row = target_table.add_row()
                        
                        # Isi data ke baris yang sudah disediakan/dibuat
                        set_cell_text_with_font(target_row.cells[0], f"{index + 1}.")
                        set_cell_text_with_font(target_row.cells[1], str(row_data["Nama/Jabatan"]))
                        set_cell_text_with_font(target_row.cells[2], str(row_data["NIP/Golongan"]))
                        set_cell_text_with_font(target_row.cells[3], str(row_data["Posisi"]))
                        set_cell_text_with_font(target_row.cells[4], f"Rp{row_data['Honor']}")
                    
                    # --- PERUBAHAN BARU: TAMBAH 1 BARIS KOSONG TERAKHIR SEBAGAI SPASI SEBELUM GARIS ---
                    final_spacer = target_table.add_row()
                    for cell in final_spacer.cells:
                        set_bottom_border(cell)
                            
                else:
                    st.warning("⚠️ Peringatan: Teks {nama} tidak ditemukan di dalam tabel Word.")

                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.session_state.doc_data = doc_io.getvalue()
                st.session_state.doc_name = f"SK_{tentang.replace(' ', '_')}.docx"
                st.session_state.doc_ready = True
                
        except Exception as e:
            st.error(f"Terjadi kesalahan teknis: {e}")
            st.error(traceback.format_exc())
            st.session_state.doc_ready = False

# 4. Menampilkan Tombol Unduh
if st.session_state.doc_ready:
    st.success("✅ Dokumen berhasil diproses (Ditambah spasi penutup sebelum garis ganda)!")
    st.download_button(
        label="⬇️ Unduh Dokumen Hasil",
        data=st.session_state.doc_data,
        file_name=st.session_state.doc_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
